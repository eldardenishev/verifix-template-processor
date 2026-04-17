from __future__ import annotations
import json
import shutil
import tempfile
import uuid
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles

from app.config import SOURCES_PATH, TEMP_DIR
from app.models import Source, AnalysisResult, MappingEntry, UnmappedEntry
from app.services.doc_converter import convert_doc_to_docx
from app.services.extractor import extract_paragraphs, get_full_text
from app.services.entity_recognizer import recognize_entities
from app.services.source_matcher import match_source
from app.services.mapper import map_entities_to_variables, filter_dynamic_mappings
from app.services.docx_writer import apply_replacements

from docx import Document

app = FastAPI(
    title="Verifix Template Processor",
    description="Microservice for converting filled Word documents into Verifix MERGEFIELD templates",
    version="0.1.0",
)

STATIC_DIR = Path(__file__).parent / "static"


@app.get("/", response_class=HTMLResponse)
async def root():
    """Serve the web UI."""
    index_path = STATIC_DIR / "index.html"
    return HTMLResponse(content=index_path.read_text(encoding="utf-8"))


def _load_sources() -> list[Source]:
    """Load sources from sources.json."""
    with open(SOURCES_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [Source(**s) for s in data["sources"]]


def _save_upload(upload: UploadFile) -> Path:
    """Save uploaded file to a temporary location."""
    suffix = Path(upload.filename or "document.docx").suffix
    tmp_path = TEMP_DIR / f"{uuid.uuid4()}{suffix}"
    with open(tmp_path, "wb") as f:
        shutil.copyfileobj(upload.file, f)
    return tmp_path


def _run_pipeline(file_path: Path) -> tuple[Document, AnalysisResult]:
    """Full processing pipeline: extract -> recognize -> match -> map -> write."""
    # Step 1: Convert .doc to .docx if needed
    if file_path.suffix.lower() == ".doc":
        file_path = convert_doc_to_docx(file_path)

    # Step 2: Load document and extract paragraphs
    doc = Document(str(file_path))
    paragraphs = extract_paragraphs(doc)
    full_text = get_full_text(paragraphs)

    # Step 3: Recognize entities
    entities = recognize_entities(full_text)

    # Step 4: Load sources and match
    sources = _load_sources()
    best_source, confidence = match_source(entities, sources)

    # Step 5: Map entities to variables
    mappings, unmapped = map_entities_to_variables(entities, best_source)

    # Step 6: Filter dynamic vs static
    dynamic_mappings, static_mappings = filter_dynamic_mappings(mappings)

    # Build analysis result
    analysis = AnalysisResult(
        detected_source=best_source.name,
        detected_source_id=best_source.id,
        confidence=confidence,
        mappings=mappings,
        unmapped=unmapped,
    )

    # Step 7: Apply replacements to document
    # Re-load document fresh (since extract_paragraphs doesn't modify it)
    doc = Document(str(file_path))
    doc = apply_replacements(doc, dynamic_mappings, unmapped, analysis)

    return doc, analysis


# ───────────────────────────────────────────
# Endpoints
# ───────────────────────────────────────────

@app.post("/api/v1/process")
async def process_document(file: UploadFile = File(...)):
    """Process a document and return a .docx with MERGEFIELDs."""
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    suffix = Path(file.filename).suffix.lower()
    if suffix not in (".doc", ".docx"):
        raise HTTPException(400, f"Unsupported file type: {suffix}. Use .doc or .docx")

    tmp_path = _save_upload(file)
    try:
        doc, analysis = _run_pipeline(tmp_path)

        # Save result
        output_path = TEMP_DIR / f"result_{uuid.uuid4()}.docx"
        doc.save(str(output_path))

        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"template_{Path(file.filename).stem}.docx",
            headers={
                "X-Detected-Source": analysis.detected_source,
                "X-Confidence": f"{analysis.confidence:.2f}",
                "X-Mappings-Count": str(len([m for m in analysis.mappings if m.dynamic])),
                "X-Unmapped-Count": str(len(analysis.unmapped)),
            },
        )
    except Exception as e:
        raise HTTPException(500, detail=str(e))
    finally:
        # Cleanup input file
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


@app.post("/api/v1/analyze")
async def analyze_document(file: UploadFile = File(...)):
    """Analyze a document and return JSON with mapping details (no file generation)."""
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    suffix = Path(file.filename).suffix.lower()
    if suffix not in (".doc", ".docx"):
        raise HTTPException(400, f"Unsupported file type: {suffix}. Use .doc or .docx")

    tmp_path = _save_upload(file)
    try:
        # Run pipeline without docx_writer
        if tmp_path.suffix.lower() == ".doc":
            tmp_path = convert_doc_to_docx(tmp_path)

        doc = Document(str(tmp_path))
        paragraphs = extract_paragraphs(doc)
        full_text = get_full_text(paragraphs)
        entities = recognize_entities(full_text)
        sources = _load_sources()
        best_source, confidence = match_source(entities, sources)
        mappings, unmapped = map_entities_to_variables(entities, best_source)

        result = AnalysisResult(
            detected_source=best_source.name,
            detected_source_id=best_source.id,
            confidence=confidence,
            mappings=mappings,
            unmapped=unmapped,
        )

        return JSONResponse(content=result.model_dump())
    except Exception as e:
        raise HTTPException(500, detail=str(e))
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


@app.get("/api/v1/sources")
async def list_sources():
    """List all available Verifix sources with their variables."""
    sources = _load_sources()
    return [s.model_dump() for s in sources]


@app.get("/health")
async def health():
    """Health check endpoint."""
    return {"status": "ok", "version": "0.1.0"}
