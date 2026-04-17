from __future__ import annotations
import copy
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from app.models import MappingEntry, UnmappedEntry, AnalysisResult
from app.utils.mergefield import make_mergefield_with_format


def apply_replacements(
    doc: Document,
    dynamic_mappings: list[MappingEntry],
    unmapped: list[UnmappedEntry],
    analysis: AnalysisResult,
) -> Document:
    """
    Main function: replaces dynamic entity values with MERGEFIELDs,
    highlights unmapped entities in red, and adds a summary block.
    """
    # Build replacement map: original_text -> variable_code
    replace_map: dict[str, str] = {}
    for m in dynamic_mappings:
        replace_map[m.original] = m.variable

    # Build unmapped set for red highlighting
    unmapped_texts: set[str] = {u.original for u in unmapped}

    # Process body paragraphs
    for para in doc.paragraphs:
        _process_paragraph(para, replace_map, unmapped_texts)

    # Process table cells
    for table in doc.tables:
        seen_in_row: dict[int, set[str]] = {}
        for r_idx, row in enumerate(table.rows):
            seen_in_row[r_idx] = set()
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text
                    if not para_text.strip():
                        continue
                    # Handle merged cells: only process first occurrence per row
                    if para_text in seen_in_row[r_idx]:
                        continue
                    seen_in_row[r_idx].add(para_text)
                    _process_paragraph(para, replace_map, unmapped_texts)

    # Add summary block at the beginning
    _add_summary_block(doc, analysis)

    # Add unmapped section at the end
    if unmapped:
        _add_unmapped_section(doc, unmapped)

    return doc


def _process_paragraph(
    para: Paragraph,
    replace_map: dict[str, str],
    unmapped_texts: set[str],
) -> None:
    """
    Process a single paragraph: replace mapped values with MERGEFIELDs
    and highlight unmapped values in red.
    """
    full_text = para.text
    if not full_text.strip():
        return

    # Check if any replacement or highlight is needed
    needs_processing = False
    for original in replace_map:
        if original in full_text:
            needs_processing = True
            break
    if not needs_processing:
        for original in unmapped_texts:
            if original in full_text:
                needs_processing = True
                break
    if not needs_processing:
        return

    # Collect all replacements/highlights with their positions
    actions: list[dict] = []

    for original, variable in replace_map.items():
        start = 0
        while True:
            pos = full_text.find(original, start)
            if pos == -1:
                break
            actions.append({
                "start": pos,
                "end": pos + len(original),
                "type": "mergefield",
                "variable": variable,
                "original": original,
            })
            start = pos + len(original)

    for original in unmapped_texts:
        start = 0
        while True:
            pos = full_text.find(original, start)
            if pos == -1:
                break
            # Check this position isn't already covered by a mergefield action
            covered = False
            for a in actions:
                if a["start"] <= pos < a["end"]:
                    covered = True
                    break
            if not covered:
                actions.append({
                    "start": pos,
                    "end": pos + len(original),
                    "type": "highlight_red",
                    "original": original,
                })
            start = pos + len(original)

    if not actions:
        return

    # Sort actions by position
    actions.sort(key=lambda a: a["start"])

    # Remove overlapping actions (keep first)
    filtered_actions: list[dict] = []
    for action in actions:
        overlap = False
        for existing in filtered_actions:
            if action["start"] < existing["end"] and action["end"] > existing["start"]:
                overlap = True
                break
        if not overlap:
            filtered_actions.append(action)

    # Get formatting from the first run (as default)
    default_rpr = None
    if para.runs:
        first_run_elem = para.runs[0]._element
        rpr = first_run_elem.find(qn('w:rPr'))
        if rpr is not None:
            default_rpr = copy.deepcopy(rpr)

    # Rebuild paragraph: clear existing runs, insert new content
    _rebuild_paragraph(para, full_text, filtered_actions, default_rpr)


def _rebuild_paragraph(
    para: Paragraph,
    full_text: str,
    actions: list[dict],
    default_rpr: OxmlElement | None,
) -> None:
    """
    Clears the paragraph's runs and rebuilds it with:
    - Normal text runs for unchanged parts
    - MERGEFIELD elements for replacements
    - Red-colored runs for unmapped entities
    """
    p_elem = para._element

    # Remove all existing runs and field elements from paragraph
    for child in list(p_elem):
        tag = child.tag
        if tag.endswith('}r') or tag.endswith('}fldSimple') or tag.endswith('}hyperlink'):
            p_elem.remove(child)

    # Build new content
    pos = 0
    for action in actions:
        # Text before this action
        if pos < action["start"]:
            before_text = full_text[pos:action["start"]]
            if before_text:
                _add_text_run(p_elem, before_text, default_rpr)

        if action["type"] == "mergefield":
            # Insert MERGEFIELD
            fld = make_mergefield_with_format(action["variable"], _make_run_with_rpr(default_rpr))
            p_elem.append(fld)
        elif action["type"] == "highlight_red":
            # Insert red-colored text
            _add_text_run(p_elem, action["original"], default_rpr, color="FF0000")

        pos = action["end"]

    # Text after last action
    if pos < len(full_text):
        remaining = full_text[pos:]
        if remaining:
            _add_text_run(p_elem, remaining, default_rpr)


def _add_text_run(
    p_elem: OxmlElement,
    text: str,
    rpr_template: OxmlElement | None,
    color: str | None = None,
) -> None:
    """Adds a text run to the paragraph element, preserving formatting."""
    r = OxmlElement('w:r')

    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
        if color:
            # Remove existing color if any
            existing_color = rpr.find(qn('w:color'))
            if existing_color is not None:
                rpr.remove(existing_color)
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), color)
            rpr.append(color_elem)
        r.append(rpr)
    elif color:
        rpr = OxmlElement('w:rPr')
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), color)
        rpr.append(color_elem)
        r.append(rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)


def _make_run_with_rpr(rpr_template: OxmlElement | None) -> OxmlElement | None:
    """Creates a minimal run element with formatting for MERGEFIELD creation."""
    if rpr_template is None:
        return None
    r = OxmlElement('w:r')
    r.append(copy.deepcopy(rpr_template))
    return r


def _add_summary_block(doc: Document, analysis: AnalysisResult) -> None:
    """Adds a summary block at the beginning of the document."""
    # Insert before the first paragraph
    body = doc.element.body

    # Create summary paragraphs
    summary_lines = [
        f"=== VERIFIX TEMPLATE PROCESSOR ===",
        f"Источник: {analysis.detected_source} (ID: {analysis.detected_source_id})",
        f"Уверенность: {analysis.confidence:.0%}",
        f"Замен MERGEFIELD: {len([m for m in analysis.mappings if m.dynamic])}",
        f"Не сопоставлено: {len(analysis.unmapped)}",
        f"=================================",
    ]

    # Insert summary paragraphs at the beginning (in reverse order)
    first_child = body[0] if len(body) > 0 else None

    for line in reversed(summary_lines):
        p = OxmlElement('w:p')

        # Add paragraph properties for small gray font
        ppr = OxmlElement('w:pPr')
        p.append(ppr)

        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')

        # Small font
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '16')  # 8pt
        rpr.append(sz)

        # Gray color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rpr.append(color)

        r.append(rpr)

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
        p.append(r)

        if first_child is not None:
            body.insert(body.index(first_child), p)
        else:
            body.append(p)


def _add_unmapped_section(doc: Document, unmapped: list[UnmappedEntry]) -> None:
    """Adds a section at the end listing unmapped values."""
    # Separator
    doc.add_paragraph("")
    sep = doc.add_paragraph("=" * 50)
    for run in sep.runs:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run.font.size = Pt(8)

    header = doc.add_paragraph("ПРОВЕРЬТЕ ВРУЧНУЮ (не сопоставленные значения):")
    for run in header.runs:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run.bold = True
        run.font.size = Pt(9)

    for entry in unmapped:
        line = f"  - \"{entry.original}\" [{entry.entity_type}] — {entry.reason}"
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(8)
