from __future__ import annotations
from dataclasses import dataclass, field
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class ParagraphInfo:
    index: int
    text: str
    runs: list[RunDetail] = field(default_factory=list)
    location: str = "body"  # "body" or "table"
    table_index: int = -1
    row_index: int = -1
    cell_index: int = -1


@dataclass
class RunDetail:
    text: str
    start: int  # start offset within paragraph text
    end: int
    bold: bool | None = None
    italic: bool | None = None
    font_name: str | None = None
    font_size: float | None = None


def extract_paragraphs(doc: Document) -> list[ParagraphInfo]:
    """Extracts all paragraphs from the document body and tables."""
    paragraphs: list[ParagraphInfo] = []
    idx = 0

    # Body paragraphs
    for para in doc.paragraphs:
        pinfo = _extract_paragraph(para, idx, location="body")
        if pinfo.text.strip():
            paragraphs.append(pinfo)
        idx += 1

    # Table paragraphs
    for t_idx, table in enumerate(doc.tables):
        seen_in_row: dict[int, set[str]] = {}
        for r_idx, row in enumerate(table.rows):
            seen_in_row[r_idx] = set()
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    cell_text = para.text.strip()
                    if not cell_text:
                        continue
                    # Skip duplicate text in merged cells (same row)
                    if cell_text in seen_in_row[r_idx]:
                        continue
                    seen_in_row[r_idx].add(cell_text)

                    pinfo = _extract_paragraph(
                        para, idx,
                        location="table",
                        table_index=t_idx,
                        row_index=r_idx,
                        cell_index=c_idx,
                    )
                    paragraphs.append(pinfo)
                    idx += 1

    return paragraphs


def _extract_paragraph(
    para: Paragraph,
    index: int,
    location: str = "body",
    table_index: int = -1,
    row_index: int = -1,
    cell_index: int = -1,
) -> ParagraphInfo:
    """Extracts text and run details from a single paragraph."""
    runs: list[RunDetail] = []
    offset = 0
    for run in para.runs:
        text = run.text
        rd = RunDetail(
            text=text,
            start=offset,
            end=offset + len(text),
            bold=run.bold,
            italic=run.italic,
            font_name=run.font.name if run.font else None,
            font_size=run.font.size.pt if run.font and run.font.size else None,
        )
        runs.append(rd)
        offset += len(text)

    return ParagraphInfo(
        index=index,
        text=para.text,
        runs=runs,
        location=location,
        table_index=table_index,
        row_index=row_index,
        cell_index=cell_index,
    )


def get_full_text(paragraphs: list[ParagraphInfo]) -> str:
    """Joins all paragraph texts with newlines for entity recognition."""
    return "\n".join(p.text for p in paragraphs)
